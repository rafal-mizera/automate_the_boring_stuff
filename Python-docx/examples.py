import docx
import read_docx
from docx.text.run import *

doc = docx.Document("demo.docx")

# print(len(doc.paragraphs))
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)
#
# print(doc.paragraphs[1].runs)
# print(doc.paragraphs[1].runs[0].text)
#
# print(read_docx.getText("demo.docx"))
# print(doc.paragraphs[0].style)
# doc.paragraphs[0].style = "Normal"
#
# doc.paragraphs[1].style = "Quote"
# doc.paragraphs[2].underline = True
# doc.save("restyled.docx")

doc = docx.Document()
doc.add_paragraph("Hello World","Title")
doc.save("helloworld.docx")
para_obj = doc.add_paragraph("123")
para_obj.add_run("456")
doc.save("helloworld.docx")

for x in range(5):
    doc.add_heading(f"Header{x}",x)

doc.save("helloworld.docx")

par = doc.add_paragraph("This is on the first page")
run = par.add_run()
run.add_break(WD_BREAK.PAGE)

doc.add_paragraph("This is on the second page")
doc.save("helloworld.docx")
